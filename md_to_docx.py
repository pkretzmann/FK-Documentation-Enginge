#!/usr/bin/env python3
"""
Markdown to DOCX Converter
==========================
Konverterer Markdown-filer til professionelle Word-dokumenter med
Fich & Kretzmann branding (turkis/teal og mørkeblå farver).

Funktioner:
- Headers med stilfuld formatering
- Kodeblokke med monospace font og baggrund
- Tabeller med alternerende rækkefarver
- Billeder med caption
- Lister (punktopstilling og nummererede)
- Logo i header
- Moderne, professionelt design

Brug:
    python md_to_docx.py input.md output.docx
    python md_to_docx.py input.md  # Output bliver input.docx
"""

import re
import sys
import os
import tempfile
import urllib.request
from pathlib import Path
from typing import Optional, List, Tuple
import argparse

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# =============================================================================
# FARVEPALETTE - Fich & Kretzmann Branding
# =============================================================================
class Colors:
    """Fich & Kretzmann farvepalette baseret på logo"""
    # Primære farver fra logo
    TEAL = RGBColor(0, 165, 168)           # #00A5A8 - Turkis/teal
    DARK_BLUE = RGBColor(27, 58, 95)       # #1B3A5F - Mørkeblå

    # Sekundære farver
    LIGHT_TEAL = RGBColor(230, 247, 247)   # Lys turkis baggrund
    MEDIUM_TEAL = RGBColor(0, 139, 142)    # Mørkere teal til accenter
    LIGHT_BLUE = RGBColor(240, 244, 248)   # Lys blå baggrund

    # Neutrale farver
    WHITE = RGBColor(255, 255, 255)
    LIGHT_GRAY = RGBColor(245, 245, 245)   # Baggrund til kodeblokke
    MEDIUM_GRAY = RGBColor(128, 128, 128)
    DARK_GRAY = RGBColor(51, 51, 51)       # Tekst
    CODE_BG = RGBColor(248, 249, 250)      # Kodeblok baggrund
    CODE_BORDER = RGBColor(225, 228, 232)  # Kodeblok kant


# =============================================================================
# SKRIFTTYPER
# =============================================================================
class Fonts:
    """Professionelle skrifttyper"""
    HEADING = "Segoe UI"           # Moderne sans-serif til overskrifter
    BODY = "Calibri"               # Læsbar brødtekst
    CODE = "Cascadia Code"         # Monospace til kode (fallback: Consolas)
    CODE_FALLBACK = "Consolas"


# =============================================================================
# STØRRELSER
# =============================================================================
class Sizes:
    """Skriftstørrelser i Pt"""
    H1 = 28
    H2 = 22
    H3 = 16
    H4 = 14
    H5 = 12
    H6 = 11
    BODY = 11
    CODE = 10
    SMALL = 9
    FOOTER = 9


# =============================================================================
# DOKUMENT BUILDER
# =============================================================================
class DocumentBuilder:
    """Bygger Word-dokumenter fra Markdown"""

    def __init__(self, logo_path: Optional[str] = None):
        self.doc = Document()
        self.logo_path = logo_path
        self._setup_styles()
        self._setup_page()

    def _setup_page(self):
        """Konfigurer sidemargin og layout"""
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

            # Header og footer afstand
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)

    def _setup_styles(self):
        """Opret brugerdefinerede stilarter"""
        styles = self.doc.styles

        # Heading 1 - Stor, fed, mørkeblå
        self._create_heading_style(styles, 'Heading 1', Sizes.H1, Colors.DARK_BLUE,
                                   space_before=18, space_after=12, bold=True)

        # Heading 2 - Teal farve
        self._create_heading_style(styles, 'Heading 2', Sizes.H2, Colors.TEAL,
                                   space_before=16, space_after=8, bold=True)

        # Heading 3 - Mørkeblå
        self._create_heading_style(styles, 'Heading 3', Sizes.H3, Colors.DARK_BLUE,
                                   space_before=14, space_after=6, bold=True)

        # Heading 4
        self._create_heading_style(styles, 'Heading 4', Sizes.H4, Colors.MEDIUM_TEAL,
                                   space_before=12, space_after=4, bold=True)

        # Heading 5 & 6
        self._create_heading_style(styles, 'Heading 5', Sizes.H5, Colors.DARK_GRAY,
                                   space_before=10, space_after=4, bold=True)
        self._create_heading_style(styles, 'Heading 6', Sizes.H6, Colors.MEDIUM_GRAY,
                                   space_before=10, space_after=4, bold=True, italic=True)

        # Normal tekst
        normal = styles['Normal']
        normal.font.name = Fonts.BODY
        normal.font.size = Pt(Sizes.BODY)
        normal.font.color.rgb = Colors.DARK_GRAY
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15

        # Kodeblok stil
        if 'Code Block' not in [s.name for s in styles]:
            code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = Fonts.CODE
            code_style.font.size = Pt(Sizes.CODE)
            code_style.font.color.rgb = Colors.DARK_GRAY
            code_style.paragraph_format.space_before = Pt(8)
            code_style.paragraph_format.space_after = Pt(8)
            code_style.paragraph_format.left_indent = Cm(0.5)

        # Inline kode
        if 'Code Char' not in [s.name for s in styles]:
            code_char = styles.add_style('Code Char', WD_STYLE_TYPE.CHARACTER)
            code_char.font.name = Fonts.CODE
            code_char.font.size = Pt(Sizes.CODE)
            code_char.font.color.rgb = Colors.TEAL

        # Citat/blockquote
        if 'Quote' not in [s.name for s in styles]:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.font.name = Fonts.BODY
            quote_style.font.size = Pt(Sizes.BODY)
            quote_style.font.italic = True
            quote_style.font.color.rgb = Colors.MEDIUM_GRAY
            quote_style.paragraph_format.left_indent = Cm(1)
            quote_style.paragraph_format.space_before = Pt(8)
            quote_style.paragraph_format.space_after = Pt(8)

    def _create_heading_style(self, styles, name: str, size: int, color: RGBColor,
                              space_before: int = 12, space_after: int = 6,
                              bold: bool = True, italic: bool = False):
        """Hjælpefunktion til at konfigurere heading-stilarter"""
        style = styles[name]
        style.font.name = Fonts.HEADING
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True

    def add_header_with_logo(self, title: Optional[str] = None):
        """Tilføj header med logo"""
        section = self.doc.sections[0]
        header = section.header

        # Opret tabel til header layout (logo til venstre, evt. titel til højre)
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.autofit = False
        header_table.allow_autofit = False

        # Logo celle
        logo_cell = header_table.cell(0, 0)
        logo_cell.width = Inches(1.5)

        if self.logo_path and os.path.exists(self.logo_path):
            logo_para = logo_cell.paragraphs[0]
            logo_para.paragraph_format.space_after = Pt(0)
            logo_run = logo_para.add_run()
            logo_run.add_picture(self.logo_path, width=Inches(0.8))

        # Titel celle (højrejusteret)
        title_cell = header_table.cell(0, 1)
        title_cell.width = Inches(5)
        if title:
            title_para = title_cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run(title)
            title_run.font.name = Fonts.HEADING
            title_run.font.size = Pt(Sizes.SMALL)
            title_run.font.color.rgb = Colors.MEDIUM_GRAY

        # Tilføj linje under header
        self._add_header_line(section)

    def _add_header_line(self, section):
        """Tilføj en farvet linje under header"""
        header = section.header
        line_para = header.add_paragraph()
        line_para.paragraph_format.space_before = Pt(0)
        line_para.paragraph_format.space_after = Pt(0)

        # Tilføj en "linje" ved at bruge border
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # Linjetykkelse
        bottom.set(qn('w:color'), '00A5A8')  # Teal farve
        pBdr.append(bottom)
        line_para._p.get_or_add_pPr().append(pBdr)

    def add_footer(self, text: Optional[str] = None):
        """Tilføj footer med sidenummer"""
        section = self.doc.sections[0]
        footer = section.footer

        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Tilføj linje over footer
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:color'), 'CCCCCC')
        pBdr.append(top)
        footer_para._p.get_or_add_pPr().append(pBdr)

        # Footer tekst
        if text:
            run = footer_para.add_run(f"{text}  |  ")
            run.font.size = Pt(Sizes.FOOTER)
            run.font.color.rgb = Colors.MEDIUM_GRAY

        # Sidenummer
        run = footer_para.add_run("Side ")
        run.font.size = Pt(Sizes.FOOTER)
        run.font.color.rgb = Colors.MEDIUM_GRAY

        # Indsæt sidenummer felt
        self._add_page_number(footer_para)

    def _add_page_number(self, paragraph):
        """Indsæt sidenummer felt"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(Sizes.FOOTER)
        run.font.color.rgb = Colors.MEDIUM_GRAY

    def add_title_page(self, title: str, subtitle: Optional[str] = None,
                       author: Optional[str] = None, date: Optional[str] = None):
        """Opret en flot titelside"""
        # Logo centreret øverst
        if self.logo_path and os.path.exists(self.logo_path):
            logo_para = self.doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.paragraph_format.space_before = Pt(60)
            run = logo_para.add_run()
            run.add_picture(self.logo_path, width=Inches(2))

        # Titel
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(80)
        title_run = title_para.add_run(title)
        title_run.font.name = Fonts.HEADING
        title_run.font.size = Pt(36)
        title_run.font.color.rgb = Colors.DARK_BLUE
        title_run.font.bold = True

        # Dekorativ linje
        line_para = self.doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("─" * 30)
        line_run.font.color.rgb = Colors.TEAL
        line_run.font.size = Pt(14)

        # Undertitel
        if subtitle:
            sub_para = self.doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.paragraph_format.space_before = Pt(20)
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = Fonts.HEADING
            sub_run.font.size = Pt(18)
            sub_run.font.color.rgb = Colors.TEAL

        # Forfatter og dato
        if author or date:
            info_para = self.doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_para.paragraph_format.space_before = Pt(100)

            if author:
                author_run = info_para.add_run(f"{author}\n")
                author_run.font.name = Fonts.BODY
                author_run.font.size = Pt(12)
                author_run.font.color.rgb = Colors.DARK_GRAY

            if date:
                date_run = info_para.add_run(date)
                date_run.font.name = Fonts.BODY
                date_run.font.size = Pt(11)
                date_run.font.color.rgb = Colors.MEDIUM_GRAY

        # Sideskift efter titelside
        self.doc.add_page_break()

    def add_heading(self, text: str, level: int = 1):
        """Tilføj overskrift"""
        heading = self.doc.add_heading(text, level=min(level, 6))
        return heading

    def add_paragraph(self, text: str, style: str = 'Normal') -> 'Paragraph':
        """Tilføj afsnit med inline formatering"""
        para = self.doc.add_paragraph(style=style)
        self._parse_inline_formatting(para, text)
        return para

    def _parse_inline_formatting(self, paragraph, text: str):
        """Parser og tilføjer inline formatering (fed, kursiv, kode, links)"""
        # Regex patterns for inline formatting
        patterns = [
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),      # ***bold italic***
            (r'\*\*(.+?)\*\*', 'bold'),                  # **bold**
            (r'__(.+?)__', 'bold'),                      # __bold__
            (r'\*(.+?)\*', 'italic'),                    # *italic*
            (r'_(.+?)_', 'italic'),                      # _italic_
            (r'`(.+?)`', 'code'),                        # `code`
            (r'\[(.+?)\]\((.+?)\)', 'link'),            # [text](url)
            (r'~~(.+?)~~', 'strikethrough'),            # ~~strikethrough~~
        ]

        # Simpel implementation - kombiner alle patterns
        combined_pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|__.+?__|_[^_]+?_|\*.+?\*|`.+?`|\[.+?\]\(.+?\)|~~.+?~~)'

        parts = re.split(combined_pattern, text)

        for part in parts:
            if not part:
                continue

            # Check hver type formatering
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = Fonts.CODE
                run.font.size = Pt(Sizes.CODE)
                run.font.color.rgb = Colors.TEAL
            elif part.startswith('~~') and part.endswith('~~'):
                run = paragraph.add_run(part[2:-2])
                run.font.strike = True
            elif part.startswith('[') and '](' in part:
                # Link format: [text](url)
                match = re.match(r'\[(.+?)\]\((.+?)\)', part)
                if match:
                    link_text, url = match.groups()
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = Colors.TEAL
                    run.font.underline = True
                    # Note: python-docx understøtter ikke hyperlinks direkte i runs
            else:
                run = paragraph.add_run(part)

    def add_code_block(self, code: str, language: Optional[str] = None):
        """Tilføj kodeblok med baggrund"""
        # Tilføj en tabel med én celle for at simulere baggrund
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False

        cell = table.cell(0, 0)

        # Sæt baggrundsfarve på cellen
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F8F9FA')  # Lys grå baggrund
        cell._tc.get_or_add_tcPr().append(shading)

        # Tilføj kant
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'E1E4E8')
            tcBorders.append(border)
        cell._tc.get_or_add_tcPr().append(tcBorders)

        # Tilføj kodelinjer
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)

        # Tilføj sprog-label hvis angivet
        if language:
            lang_run = para.add_run(f"{language}\n")
            lang_run.font.name = Fonts.CODE
            lang_run.font.size = Pt(Sizes.SMALL)
            lang_run.font.color.rgb = Colors.MEDIUM_GRAY
            lang_run.font.bold = True

        # Tilføj koden
        code_run = para.add_run(code)
        code_run.font.name = Fonts.CODE
        code_run.font.size = Pt(Sizes.CODE)
        code_run.font.color.rgb = Colors.DARK_GRAY

        # Tilføj lidt mellemrum efter kodeblokken
        self.doc.add_paragraph()

    def add_blockquote(self, text: str):
        """Tilføj citat/blockquote"""
        # Opret tabel for at få venstre kant
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Venstre kant i teal farve
        tcBorders = OxmlElement('w:tcBorders')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:color'), '00A5A8')  # Teal
        tcBorders.append(left)
        cell._tc.get_or_add_tcPr().append(tcBorders)

        # Lys baggrund
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F4F8')
        cell._tc.get_or_add_tcPr().append(shading)

        # Tekst
        para = cell.paragraphs[0]
        para.paragraph_format.left_indent = Cm(0.3)
        run = para.add_run(text)
        run.font.italic = True
        run.font.color.rgb = Colors.DARK_GRAY

    def add_bullet_list(self, items: List[str], level: int = 0):
        """Tilføj punktopstilling"""
        for item in items:
            para = self.doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
            self._parse_inline_formatting(para, item)

    def add_numbered_list(self, items: List[str], level: int = 0):
        """Tilføj nummereret liste"""
        for item in items:
            para = self.doc.add_paragraph(style='List Number')
            para.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
            self._parse_inline_formatting(para, item)

    def add_table(self, headers: List[str], rows: List[List[str]]):
        """Tilføj tabel med Fich & Kretzmann styling"""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header række
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.bold = True
            run.font.color.rgb = Colors.WHITE
            run.font.name = Fonts.HEADING
            run.font.size = Pt(Sizes.BODY)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Mørkeblå baggrund til header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1B3A5F')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rækker med alternerende farver
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                para = cell.paragraphs[0]
                # Parse inline formatting (bold, italic, etc.) in table cells
                self._parse_inline_formatting(para, str(cell_text))
                # Apply default font settings to all runs in the cell
                for run in para.runs:
                    if run.font.name is None:
                        run.font.name = Fonts.BODY
                    if run.font.size is None:
                        run.font.size = Pt(Sizes.BODY)
                    if run.font.color.rgb is None:
                        run.font.color.rgb = Colors.DARK_GRAY

                # Alternerende baggrundsfarve
                if row_idx % 2 == 1:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E6F7F7')  # Lys teal
                    cell._tc.get_or_add_tcPr().append(shading)

        # Tilføj mellemrum efter tabel
        self.doc.add_paragraph()

    def add_horizontal_rule(self):
        """Tilføj horisontal linje"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '00A5A8')
        pBdr.append(bottom)
        para._p.get_or_add_pPr().append(pBdr)

    def add_image(self, image_path: str, caption: Optional[str] = None,
                  width: Optional[float] = None):
        """Tilføj billede med optional caption"""
        if not os.path.exists(image_path):
            # Tilføj placeholder tekst hvis billede ikke findes
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[Billede ikke fundet: {image_path}]")
            run.font.italic = True
            run.font.color.rgb = Colors.MEDIUM_GRAY
            return

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()

        if width:
            run.add_picture(image_path, width=Inches(width))
        else:
            run.add_picture(image_path, width=Inches(5))  # Default bredde

        # Tilføj caption
        if caption:
            caption_para = self.doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            caption_run.font.italic = True
            caption_run.font.size = Pt(Sizes.SMALL)
            caption_run.font.color.rgb = Colors.MEDIUM_GRAY

    def save(self, output_path: str):
        """Gem dokumentet"""
        self.doc.save(output_path)
        print(f"Dokument gemt: {output_path}")


# =============================================================================
# MARKDOWN PARSER
# =============================================================================
class MarkdownParser:
    """Parser Markdown og konverterer til DocumentBuilder kald"""

    def __init__(self, builder: DocumentBuilder, base_path: str = "."):
        self.builder = builder
        self.base_path = base_path

    def parse_file(self, filepath: str):
        """Parse en Markdown fil"""
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        self.base_path = os.path.dirname(os.path.abspath(filepath))
        self.parse(content)

    def parse(self, content: str):
        """Parse Markdown indhold"""
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Kodeblok (```)
            if line.strip().startswith('```'):
                language = line.strip()[3:].strip() or None
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self.builder.add_code_block('\n'.join(code_lines), language)
                i += 1
                continue

            # Tabel
            if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                headers, rows, new_i = self._parse_table(lines, i)
                if headers:
                    self.builder.add_table(headers, rows)
                i = new_i
                continue

            # Overskrift
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if text:
                    self.builder.add_heading(text, level)
                i += 1
                continue

            # Page break (HTML div eller === )
            if 'page-break-after' in line or re.match(r'^={3,}\s*$', line.strip()):
                self.builder.doc.add_page_break()
                i += 1
                continue

            # Horisontal linje
            if re.match(r'^[\-\*_]{3,}\s*$', line.strip()):
                self.builder.add_horizontal_rule()
                i += 1
                continue

            # Blockquote
            if line.strip().startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i].strip().lstrip('>').strip())
                    i += 1
                self.builder.add_blockquote(' '.join(quote_lines))
                continue

            # Punktopstilling
            if re.match(r'^[\s]*[-*+]\s+', line):
                items = []
                while i < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[i]):
                    items.append(re.sub(r'^[\s]*[-*+]\s+', '', lines[i]))
                    i += 1
                self.builder.add_bullet_list(items)
                continue

            # Nummereret liste
            if re.match(r'^[\s]*\d+\.\s+', line):
                items = []
                while i < len(lines) and re.match(r'^[\s]*\d+\.\s+', lines[i]):
                    items.append(re.sub(r'^[\s]*\d+\.\s+', '', lines[i]))
                    i += 1
                self.builder.add_numbered_list(items)
                continue

            # Billede (håndterer både normale paths og paths med <> omkring)
            # Check for linked image first: [![alt](img_url)](link_url)
            # Collect consecutive badge images to put them on the same line
            linked_img_match = re.match(r'\[!\[(.*?)\]\(([^)]+)\)\]\(([^)]+)\)', line.strip())
            if linked_img_match:
                # Collect all consecutive linked images (badges)
                badge_images = []
                while i < len(lines):
                    match = re.match(r'\[!\[(.*?)\]\(([^)]+)\)\]\(([^)]+)\)', lines[i].strip())
                    if match:
                        alt_text, img_url, link_url = match.groups()
                        badge_images.append((img_url, alt_text))
                        i += 1
                    else:
                        break
                # Add all badges to the same paragraph
                self._add_badge_images(badge_images)
                continue

            # Standard image: ![alt](path)
            img_match = re.match(r'!\[(.*?)\]\(<?([^>]+)>?\)', line.strip())
            if img_match:
                alt_text, img_path = img_match.groups()
                self._add_image_from_url_or_path(img_path, alt_text)
                i += 1
                continue

            # Normal tekst (ikke tom linje)
            if line.strip():
                # Saml efterfølgende linjer der er del af samme afsnit
                para_lines = [line]
                i += 1
                while i < len(lines) and lines[i].strip() and \
                      not lines[i].startswith('#') and \
                      not lines[i].startswith('```') and \
                      not lines[i].startswith('>') and \
                      not re.match(r'^[\s]*[-*+]\s+', lines[i]) and \
                      not re.match(r'^[\s]*\d+\.\s+', lines[i]) and \
                      not re.match(r'^[\-\*_]{3,}\s*$', lines[i].strip()) and \
                      not ('|' in lines[i] and i + 1 < len(lines) and '---' in lines[i + 1]):
                    para_lines.append(lines[i])
                    i += 1
                self.builder.add_paragraph(' '.join(para_lines))
                continue

            i += 1

    def _add_image_from_url_or_path(self, img_source: str, alt_text: Optional[str] = None):
        """Add image from URL or local path"""
        # Check if it's a URL
        if img_source.startswith('http://') or img_source.startswith('https://'):
            # Download image to temp file
            try:
                # Handle shields.io badge URLs - they return SVG by default, need to request PNG
                if 'img.shields.io' in img_source and not img_source.endswith(('.png', '.jpg', '.jpeg', '.gif')):
                    # Append .png to get raster format instead of SVG
                    if '?' in img_source:
                        # Insert .png before query string
                        base, query = img_source.split('?', 1)
                        img_source = f"{base}.png?{query}"
                    else:
                        img_source = f"{img_source}.png"

                # Create temp file with appropriate extension
                ext = os.path.splitext(img_source.split('?')[0])[1] or '.png'
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp_path = tmp.name

                # Download with user-agent header (some servers require it)
                req = urllib.request.Request(
                    img_source,
                    headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
                )
                with urllib.request.urlopen(req, timeout=10) as response:
                    with open(tmp_path, 'wb') as f:
                        f.write(response.read())

                self.builder.add_image(tmp_path, alt_text if alt_text else None)

                # Clean up temp file
                try:
                    os.unlink(tmp_path)
                except:
                    pass
            except Exception as e:
                # If download fails, add placeholder text
                para = self.builder.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(f"[{alt_text or 'Billede'}]")
                run.font.color.rgb = Colors.TEAL
                run.font.bold = True
                print(f"Advarsel: Kunne ikke hente billede fra URL: {img_source} - {e}")
        else:
            # Local file path
            full_path = os.path.join(self.base_path, img_source)
            self.builder.add_image(full_path, alt_text if alt_text else None)

    def _add_badge_images(self, badges: List[Tuple[str, str]]):
        """Add multiple badge images on the same line with small size"""
        if not badges:
            return

        # Create a single paragraph for all badges
        para = self.builder.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)

        for idx, (img_url, alt_text) in enumerate(badges):
            # Add space between badges (except before first)
            if idx > 0:
                para.add_run("  ")  # Two spaces between badges

            try:
                # Handle shields.io badge URLs - they return SVG by default, need PNG
                if 'img.shields.io' in img_url and not img_url.endswith(('.png', '.jpg', '.jpeg', '.gif')):
                    if '?' in img_url:
                        base, query = img_url.split('?', 1)
                        img_url = f"{base}.png?{query}"
                    else:
                        img_url = f"{img_url}.png"

                # Download image
                ext = os.path.splitext(img_url.split('?')[0])[1] or '.png'
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp_path = tmp.name

                req = urllib.request.Request(
                    img_url,
                    headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
                )
                with urllib.request.urlopen(req, timeout=10) as response:
                    with open(tmp_path, 'wb') as f:
                        f.write(response.read())

                # Add image inline with small height (badges are typically 20px high)
                run = para.add_run()
                run.add_picture(tmp_path, height=Inches(0.18))  # Small badge height

                # Clean up temp file
                try:
                    os.unlink(tmp_path)
                except:
                    pass

            except Exception as e:
                # If download fails, add placeholder text
                run = para.add_run(f"[{alt_text or 'Badge'}]")
                run.font.color.rgb = Colors.TEAL
                run.font.size = Pt(Sizes.SMALL)
                print(f"Advarsel: Kunne ikke hente badge fra URL: {img_url} - {e}")

    def _parse_table(self, lines: List[str], start_idx: int) -> Tuple[List[str], List[List[str]], int]:
        """Parse Markdown tabel"""
        # Header linje - strip leading/trailing pipes, then split
        header_line = lines[start_idx].strip().strip('|')
        headers = [h.strip() for h in header_line.split('|')]

        # Spring separator linje over
        i = start_idx + 2

        # Data rækker
        rows = []
        while i < len(lines) and '|' in lines[i]:
            row_line = lines[i].strip()
            if not row_line or row_line.startswith('|---'):
                i += 1
                continue
            cells = [c.strip() for c in row_line.strip('|').split('|')]
            # Pad eller trim rækken til at matche antal headers
            if len(cells) < len(headers):
                cells.extend([''] * (len(headers) - len(cells)))
            elif len(cells) > len(headers):
                cells = cells[:len(headers)]
            rows.append(cells)
            i += 1

        return headers, rows, i


# =============================================================================
# HOVEDFUNKTION
# =============================================================================
def convert_md_to_docx(
    input_path: str,
    output_path: Optional[str] = None,
    output_folder: Optional[str] = "Docs",
    logo_path: Optional[str] = None,
    title: Optional[str] = None,
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
    include_title_page: bool = True,
    footer_text: Optional[str] = None
):
    """
    Konverterer en Markdown fil til et pænt Word dokument.

    Args:
        input_path: Sti til input .md fil
        output_path: Sti til output .docx fil (default: input filnavn med .docx)
        output_folder: Undermappe til output filer (default: Docs)
        logo_path: Sti til logo billede
        title: Dokumenttitel til titelside
        subtitle: Undertitel til titelside
        author: Forfatter til titelside
        include_title_page: Inkluder titelside (default: True)
        footer_text: Tekst i footer
    """
    # Bestem output sti
    if not output_path:
        input_dir = os.path.dirname(input_path)
        base_filename = os.path.splitext(os.path.basename(input_path))[0]

        # Opret output mappe hvis angivet
        if output_folder:
            output_dir = os.path.join(input_dir, output_folder)
            os.makedirs(output_dir, exist_ok=True)
        else:
            output_dir = input_dir

        output_path = os.path.join(output_dir, f"{base_filename}.docx")

    # Find logo automatisk hvis ikke specificeret
    if not logo_path:
        # Prøv forskellige standard placeringer
        possible_logos = [
            os.path.join(os.path.dirname(input_path), '.logo', 'FK_Logo_Square.png'),
            os.path.join(os.path.dirname(input_path), '.logo', 'FK_LOGO.jpg'),
            os.path.join(os.path.dirname(input_path), 'logo.png'),
            '.logo/FK_Logo_Square.png',
            '.logo/FK_LOGO.jpg',
        ]
        for possible_logo in possible_logos:
            if os.path.exists(possible_logo):
                logo_path = possible_logo
                break

    # Udtræk titel fra fil hvis ikke specificeret
    if not title:
        with open(input_path, 'r', encoding='utf-8') as f:
            first_lines = f.read(500)
            # Find første H1 overskrift
            h1_match = re.search(r'^#\s+(.+)$', first_lines, re.MULTILINE)
            if h1_match:
                title = h1_match.group(1).strip()
            else:
                title = os.path.splitext(os.path.basename(input_path))[0]

    # Opret dokument
    builder = DocumentBuilder(logo_path)

    # Tilføj titelside
    if include_title_page:
        from datetime import datetime
        date_str = datetime.now().strftime('%d. %B %Y')
        builder.add_title_page(title, subtitle, author, date_str)

    # Tilføj header og footer
    builder.add_header_with_logo(title)
    builder.add_footer(footer_text or "Fich & Kretzmann")

    # Parse og tilføj Markdown indhold
    parser = MarkdownParser(builder, os.path.dirname(input_path))
    parser.parse_file(input_path)

    # Gem dokument
    builder.save(output_path)

    # Konverter til PDF
    pdf_path = os.path.splitext(output_path)[0] + '.pdf'
    try:
        from docx2pdf import convert
        convert(output_path, pdf_path)
        print(f"PDF gemt: {pdf_path}")
    except ImportError:
        print("Advarsel: docx2pdf ikke installeret. Kør 'pip install docx2pdf' for PDF-understøttelse.")
    except Exception as e:
        print(f"Advarsel: Kunne ikke oprette PDF: {e}")

    return output_path


# =============================================================================
# CLI
# =============================================================================
def main():
    parser = argparse.ArgumentParser(
        description='Konverter Markdown til professionel Word dokument med Fich & Kretzmann branding',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Eksempler:
  python md_to_docx.py dokument.md
  python md_to_docx.py dokument.md output.docx
  python md_to_docx.py dokument.md -t "Min Titel" -a "Peter Kretzmann"
  python md_to_docx.py dokument.md --logo .logo/FK_Logo_Square.png --no-title-page
        """
    )

    parser.add_argument('input', help='Input Markdown fil (.md)')
    parser.add_argument('output', nargs='?', help='Output Word fil (.docx)')
    parser.add_argument('-t', '--title', help='Dokumenttitel til titelside')
    parser.add_argument('-s', '--subtitle', help='Undertitel til titelside')
    parser.add_argument('-a', '--author', help='Forfatter')
    parser.add_argument('-l', '--logo', help='Sti til logo billede')
    parser.add_argument('-f', '--footer', help='Footer tekst', default='Fich & Kretzmann')
    parser.add_argument('-o', '--output-folder', help='Undermappe til output filer', default='Docs')
    parser.add_argument('--no-output-folder', action='store_true',
                        help='Gem output i samme mappe som input')
    parser.add_argument('--no-title-page', action='store_true',
                        help='Udelad titelside')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Fejl: Filen '{args.input}' findes ikke.")
        sys.exit(1)

    try:
        output = convert_md_to_docx(
            input_path=args.input,
            output_path=args.output,
            output_folder=None if args.no_output_folder else args.output_folder,
            logo_path=args.logo,
            title=args.title,
            subtitle=args.subtitle,
            author=args.author,
            include_title_page=not args.no_title_page,
            footer_text=args.footer
        )
        print(f"Succes! Dokument oprettet: {output}")
    except Exception as e:
        print(f"Fejl under konvertering: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
